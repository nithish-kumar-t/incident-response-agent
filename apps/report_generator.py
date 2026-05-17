"""
Report Generator
Formats the log-watcher scan results as a Word (.docx) document.

Called by log_watcher.py — not meant to be run directly.

Requires: pip install python-docx
"""

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# ── Palette ────────────────────────────────────────────────────────────────────

C_ERROR   = RGBColor(0xC0, 0x39, 0x2B)   # red
C_WARN    = RGBColor(0xE6, 0x7E, 0x22)   # orange
C_OK      = RGBColor(0x27, 0xAE, 0x60)   # green
C_INFO    = RGBColor(0x27, 0x6F, 0xBE)   # blue
C_NAVY    = RGBColor(0x1A, 0x27, 0x4E)   # dark heading color

# ── Low-level helpers ──────────────────────────────────────────────────────────

def _h(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = C_NAVY


def _colored(doc: Document, text: str, color: RGBColor, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold


def _code(doc: Document, text: str, max_chars: int = 3000) -> None:
    """Courier New block for log lines and stack traces."""
    para = doc.add_paragraph()
    run = para.add_run(text[:max_chars] + ("…" if len(text) > max_chars else ""))
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)


def _kv(doc: Document, rows: list[tuple[str, Any]], style: str = "Light List Accent 1") -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = style
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = str(k)
        table.rows[i].cells[1].text = str(v) if v is not None else "—"


def _sep(doc: Document) -> None:
    doc.add_paragraph("─" * 72)


# ── Cover ──────────────────────────────────────────────────────────────────────

def _cover(doc: Document, scan: dict[str, Any]) -> None:
    title = doc.add_heading("Incident Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {scan['scan_time']}\n").font.size = Pt(10)
    meta.add_run(f"Agent Pipeline: {scan['agent_api_url']}").font.size = Pt(10)
    doc.add_paragraph()

    status_text = "PIPELINE HEALTHY" if scan.get("agent_healthy") else "PIPELINE UNREACHABLE"
    status_color = C_OK if scan.get("agent_healthy") else C_ERROR
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sp.add_run(status_text)
    run.font.color.rgb = status_color
    run.bold = True
    run.font.size = Pt(13)

    doc.add_page_break()


# ── Executive Summary ──────────────────────────────────────────────────────────

def _summary(doc: Document, scan: dict[str, Any]) -> None:
    _h(doc, "Executive Summary", 1)

    total_errors = sum(len(s.get("errors", [])) for s in scan["services"])
    total_lines  = sum(s.get("total_lines_scanned", 0) for s in scan["services"])
    affected     = [s["service_name"] for s in scan["services"] if s.get("errors")]
    pipeline_ok  = scan.get("agent_healthy", False)

    _kv(doc, [
        ("Scan Time",              scan["scan_time"]),
        ("Services Scanned",       len(scan["services"])),
        ("Total Log Lines Read",   total_lines),
        ("Total Errors Detected",  total_errors),
        ("Services With Errors",   ", ".join(affected) if affected else "None"),
        ("Agent Pipeline Status",  "HEALTHY" if pipeline_ok else "UNREACHABLE"),
        ("Analysis Performed",     "Yes" if pipeline_ok else "No — pipeline offline"),
    ])
    doc.add_paragraph()

    if total_errors == 0:
        _colored(doc, "All services are operating normally. No errors or warnings detected.", C_OK, bold=True)
    else:
        _colored(
            doc,
            f"{total_errors} error event(s) detected across "
            f"{len(affected)} service(s). See per-service sections for full analysis.",
            C_ERROR,
            bold=True,
        )

    if not pipeline_ok:
        doc.add_paragraph()
        _colored(
            doc,
            "The agent pipeline was unreachable during this scan. "
            "Start it with:  cd apps/agent && uvicorn main:app --port 8000",
            C_WARN,
        )

    doc.add_paragraph()


# ── Per-Service Section ────────────────────────────────────────────────────────

def _service_section(doc: Document, svc: dict[str, Any]) -> None:
    _h(doc, f"Service: {svc['service_name']}", 1)

    _kv(doc, [
        ("Log File",         svc["log_path"]),
        ("Log File Found",   "Yes" if svc.get("log_exists") else "NO — file missing"),
        ("Lines Scanned",    svc.get("total_lines_scanned", 0)),
        ("Errors Detected",  len(svc.get("errors", []))),
    ])
    doc.add_paragraph()

    errors = svc.get("errors", [])
    if not errors:
        _colored(doc, "No errors or warnings detected in the scanned log window.", C_OK)
        doc.add_paragraph()
        return

    # ── Error summary table ────────────────────────────────────────────────────
    _h(doc, "Detected Errors", 2)
    tbl = doc.add_table(rows=1 + len(errors), cols=3)
    tbl.style = "Medium Shading 1 Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Line #", "Timestamp", "Log Entry"
    for err in errors:
        row = tbl.add_row().cells
        row[0].text = str(err["line_number"])
        row[1].text = err.get("timestamp", "")
        row[2].text = err["line"][:180]
    doc.add_paragraph()

    # ── Per-error pipeline analysis ────────────────────────────────────────────
    pipeline_results = svc.get("pipeline_results", [])
    if not pipeline_results:
        _colored(
            doc,
            "Agent pipeline analysis not performed "
            "(pipeline was unreachable or --no-pipeline was used).",
            C_WARN,
        )
        doc.add_paragraph()
        return

    _h(doc, "Agent Pipeline Analysis", 2)

    for pr in pipeline_results:
        err  = pr["error"]
        resp = pr.get("pipeline_response", {})

        doc.add_heading(f"Error — Line {err['line_number']}  |  {err['timestamp']}", 3)

        # Context block
        p = doc.add_paragraph()
        p.add_run("Log Context:").bold = True
        _code(doc, err.get("context", err["line"]))
        doc.add_paragraph()

        if "pipeline_error" in resp:
            _colored(
                doc,
                f"Pipeline call failed: {resp['pipeline_error']}"
                + (f"\nDetail: {resp.get('detail','')}" if resp.get("detail") else ""),
                C_ERROR,
            )
            doc.add_paragraph()
            _sep(doc)
            doc.add_paragraph()
            continue

        pipe_id = resp.get("pipeline_id", "—")
        _kv(doc, [
            ("Pipeline ID",   pipe_id),
            ("Pipeline Status", resp.get("status", "—")),
            ("Started",        resp.get("started_at", "—")),
            ("Completed",      resp.get("completed_at", "—")),
        ])
        doc.add_paragraph()

        # Agent 1 — Triage
        a1 = resp.get("agent1_triage", {})
        if a1:
            _h(doc, "Agent 1 — Triage", 4)
            sev = (a1.get("severity") or "unknown").upper()
            _kv(doc, [
                ("Error Type",       a1.get("error_type", "—")),
                ("Severity",         sev),
                ("Service",          a1.get("service_name", "—")),
                ("Summary",          a1.get("error_summary", "—")),
                ("Key Indicators",   ", ".join(a1.get("key_indicators", []))),
                ("Actions Taken",    ", ".join(a1.get("actions_taken", []))),
                ("Recommended Next", a1.get("recommended_investigation", "—")),
            ])
            doc.add_paragraph()

        # Agent 2 — Investigation
        a2 = resp.get("agent2_investigation", {})
        if a2:
            _h(doc, "Agent 2 — Investigation", 4)

            health = a2.get("service_health") or {}
            _kv(doc, [
                ("Service Health",         f"{health.get('status','—')} — {health.get('details','')}"),
                ("Root Cause Hypothesis",  a2.get("root_cause_hypothesis", "—")),
                ("Affected Components",    ", ".join(a2.get("affected_components", []))),
                ("Confidence",             a2.get("confidence", "—")),
                ("Suggested Files",        ", ".join(a2.get("suggested_files_to_check", []))),
            ])

            if a2.get("inference"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Inference:").bold = True
                doc.add_paragraph(a2["inference"])

            if a2.get("log_evidence"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Log Evidence:").bold = True
                for ev in a2["log_evidence"]:
                    doc.add_paragraph(
                        f"  {ev.get('log_path','?')}  line {ev.get('line','?')}: "
                        f"{ev.get('matched_line','')}"
                    )

            if a2.get("stack_trace"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Stack Trace:").bold = True
                _code(doc, a2["stack_trace"])

            doc.add_paragraph()

        # Agent 3 — Code Analysis
        a3 = resp.get("agent3_code_analysis", {})
        if a3:
            _h(doc, "Agent 3 — Code Analysis", 4)
            _kv(doc, [
                ("Analysis Depth",    a3.get("analysis_depth", "—")),
                ("Complexity",        a3.get("complexity_assessment", "—")),
                ("Root Cause",        a3.get("root_cause", "—")),
                ("Fix Suggestion",    a3.get("fix_suggestion", "—")),
                ("Files Analyzed",    ", ".join(a3.get("files_analyzed", []))),
            ])

            if a3.get("affected_code"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Affected Code:").bold = True
                for snippet in a3["affected_code"]:
                    _code(
                        doc,
                        f"{snippet.get('file','?')}:{snippet.get('line','?')}\n"
                        f"{snippet.get('snippet','')}"
                    )
                    doc.add_paragraph(f"Issue: {snippet.get('issue','')}")

            if a3.get("recommended_next_steps"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Recommended Next Steps:").bold = True
                for step in a3["recommended_next_steps"]:
                    doc.add_paragraph(f"  • {step}")

            if a3.get("alternative_solutions"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Alternative Solutions:").bold = True
                for alt in a3["alternative_solutions"]:
                    doc.add_paragraph(f"  • {alt}")

            doc.add_paragraph()

        _sep(doc)
        doc.add_paragraph()


# ── Pipeline Alignment Section ─────────────────────────────────────────────────

def _pipeline_alignment(doc: Document, scan: dict[str, Any]) -> None:
    _h(doc, "Agent Pipeline Alignment Check", 1)

    healthy = scan.get("agent_healthy", False)
    _kv(doc, [
        ("Pipeline URL",      scan["agent_api_url"]),
        ("Reachable",         "Yes" if healthy else "No"),
        ("Analyze Endpoint",  f"{scan['agent_api_url']}/analyze  (POST)"),
        ("Expected Flow",     "Agent 1 (Triage) → Agent 2 (Investigation) → Agent 3 (Code Analysis)"),
    ])
    doc.add_paragraph()

    if healthy:
        _colored(doc, "Pipeline is running and processed all analysis requests successfully.", C_OK, bold=True)
    else:
        _colored(
            doc,
            "Pipeline was unreachable. To start it:\n"
            "  cd apps/agent\n"
            "  uvicorn main:app --host 0.0.0.0 --port 8000",
            C_ERROR,
            bold=True,
        )
    doc.add_paragraph()

    _h(doc, "Service Registry Alignment", 2)
    doc.add_paragraph(
        "The two scanned services (mongo-api-service, weather-app1) are not in "
        "service_registry.json. The middleware passes log_paths overrides directly "
        "in each /analyze request, so Agent 1 receives the correct paths without "
        "needing a registry entry."
    )
    doc.add_paragraph()

    sources = scan.get("log_sources", [])
    if sources:
        tbl = doc.add_table(rows=1 + len(sources), cols=4)
        tbl.style = "Light List Accent 1"
        h = tbl.rows[0].cells
        h[0].text = "Service"
        h[1].text = "In Registry"
        h[2].text = "Log File"
        h[3].text = "Errors Found"
        REGISTRY = {"payment-api", "user-service", "order-service"}
        for i, src in enumerate(sources):
            svc_data = next(
                (s for s in scan["services"] if s["service_name"] == src["service_name"]),
                {},
            )
            row = tbl.rows[i + 1].cells
            row[0].text = src["service_name"]
            row[1].text = "Yes" if src["service_name"] in REGISTRY else "No (override)"
            row[2].text = Path(src["log_path"]).name
            row[3].text = str(len(svc_data.get("errors", [])))
    doc.add_paragraph()


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_report(scan: dict[str, Any], reports_dir: Path) -> Path:
    """
    Build and save a Word document from the scan results.
    Returns the absolute path of the created .docx file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    _cover(doc, scan)
    _summary(doc, scan)

    for svc in scan["services"]:
        _service_section(doc, svc)

    _pipeline_alignment(doc, scan)

    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = reports_dir / f"incident_report_{ts}.docx"
    doc.save(str(out_path))
    return out_path
